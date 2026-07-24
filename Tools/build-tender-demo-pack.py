#!/usr/bin/env python3
"""Build the deterministic Tender Starter demonstration pack.

Reads reviewed plain-text sources from Tools/tender-pack/ and produces:
  1. A semantic DOCX from guide.md (en-CA, real heading styles, table headers).
  2. A PDF/UA-tagged PDF exported from the DOCX via LibreOffice Writer.
  3. An accessible XLSX from workbook.json (frozen headers, filters, wrap, widths).
  4. A deterministic ZIP with exactly three fixed-timestamp entries.

The script verifies the PDF for tagged-PDF / PDF/UA structure markers using
pypdf, normalizes PDF Info-dictionary dates, XMP timestamps, trailer /ID and
unreachable objects, normalizes XLSX container metadata, and rejects an
official-sources index that lacks HTTPS links or contains a local path.
Two isolated builds must produce byte-identical ZIPs before the committed
artifact is replaced.

Authoring-only dependencies: openpyxl, python-docx, pypdf (see requirements.txt).
"""

import hashlib
import json
import os
import re
import struct
import subprocess
import sys
import tempfile
import zipfile
from datetime import datetime, timezone
from pathlib import Path

# --- Paths ------------------------------------------------------------------

REPO_ROOT = Path(__file__).resolve().parent.parent
PACK_DIR = REPO_ROOT / "Tools" / "tender-pack"
GUIDE_MD = PACK_DIR / "guide.md"
WORKBOOK_JSON = PACK_DIR / "workbook.json"
SOURCES_TXT = PACK_DIR / "official-sources.txt"
OUTPUT_ZIP = REPO_ROOT / "Resources" / "tenders" / "tender-starter-example.zip"

DISCLAIMER = (
    "Planning aid only. Verify all requirements, deadlines, documents, and "
    "addenda at the linked official procurement source before acting or bidding."
)

# --- Fixed document instant (2026-07-24 12:00:00 UTC) -----------------------

FIXED_INSTANT = datetime(2026, 7, 24, 12, 0, 0, tzinfo=timezone.utc)
FIXED_EPOCH = int(FIXED_INSTANT.timestamp())
FIXED_PDF_DATE = "D:20260724120000+00'00'"
FIXED_XMP_DATE = b"2026-07-24T12:00:00Z"
FIXED_TIMESTAMP = (2026, 7, 24, 12, 0, 0)


def main():
    soffice = os.environ.get("TENDER_PACK_SOFFICE", "soffice")

    first = build_pack_bytes(soffice)
    second = build_pack_bytes(soffice)
    if first != second:
        raise RuntimeError(
            "Tender pack is not reproducible: two isolated builds produced different bytes"
        )
    OUTPUT_ZIP.parent.mkdir(parents=True, exist_ok=True)
    temporary_output = OUTPUT_ZIP.with_suffix(".zip.tmp")
    temporary_output.write_bytes(first)
    temporary_output.replace(OUTPUT_ZIP)
    print("Reproducibility: verified with two byte-identical isolated builds")

    # Report.
    pdf_bytes = OUTPUT_ZIP.read_bytes()
    sha = hashlib.sha256(pdf_bytes).hexdigest()
    print(f"Created: {OUTPUT_ZIP.relative_to(REPO_ROOT)}")
    print(f"Size: {len(pdf_bytes)} bytes")
    print(f"SHA-256: {sha}")
    print("PDF/UA structure: verified (StructTreeRoot, MarkInfo, Lang en-CA, pdfuaid:part)")


def build_pack_bytes(soffice: str) -> bytes:
    guide_text = GUIDE_MD.read_text(encoding="utf-8")
    workbook_data = json.loads(WORKBOOK_JSON.read_text(encoding="utf-8"))
    sources_text = SOURCES_TXT.read_text(encoding="utf-8")
    validate_sources(sources_text)

    with tempfile.TemporaryDirectory(prefix="tender-pack-build-") as work_dir:
        work = Path(work_dir)
        docx_path = work / "tender-starter-guide.docx"
        pdf_path = work / "tender-starter-guide.pdf"
        xlsx_path = work / "tender-review-workbook.xlsx"
        zip_path = work / "tender-starter-example.zip"

        build_docx(guide_text, docx_path)
        export_pdfua(docx_path, pdf_path, soffice)
        normalize_pdf(pdf_path)
        verify_pdf(pdf_path)
        build_xlsx(workbook_data, xlsx_path)
        normalize_zip_container(xlsx_path)
        create_deterministic_zip(
            zip_path,
            [
                ("tender-starter-guide.pdf", pdf_path.read_bytes()),
                ("tender-review-workbook.xlsx", xlsx_path.read_bytes()),
                ("official-sources.txt", sources_text.encode("utf-8")),
            ],
        )
        return zip_path.read_bytes()


# --- Source validation ------------------------------------------------------

def validate_sources(text: str):
    """Reject an official-sources index lacking HTTPS links or containing a local path."""
    if "https://" not in text:
        raise ValueError("official-sources.txt must contain at least one HTTPS link")
    if "file://" in text or "/Users/" in text:
        raise ValueError("official-sources.txt must not contain local file paths")


# --- DOCX -------------------------------------------------------------------

def set_style_language(doc, language: str):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    # Set the document-defaults language (this is what LibreOffice reads for
    # the PDF export catalog /Lang).
    styles_element = doc.styles.element
    doc_defaults = styles_element.find(qn("w:docDefaults"))
    if doc_defaults is None:
        doc_defaults = OxmlElement("w:docDefaults")
        styles_element.insert(0, doc_defaults)
    rpr_default = doc_defaults.find(qn("w:rPrDefault"))
    if rpr_default is None:
        rpr_default = OxmlElement("w:rPrDefault")
        doc_defaults.append(rpr_default)
    rpr = rpr_default.find(qn("w:rPr"))
    if rpr is None:
        rpr = OxmlElement("w:rPr")
        rpr_default.append(rpr)
    lang = rpr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        rpr.append(lang)
    lang.set(qn("w:val"), language)
    lang.set(qn("w:eastAsia"), language)
    lang.set(qn("w:bidi"), language)

    for style_name in (
        "Normal",
        "Title",
        "Heading 1",
        "Heading 2",
        "List Bullet",
        "List Number",
    ):
        style = doc.styles[style_name]
        run_properties = style.element.get_or_add_rPr()
        language_element = run_properties.find(qn("w:lang"))
        if language_element is None:
            language_element = OxmlElement("w:lang")
            run_properties.append(language_element)
        language_element.set(qn("w:val"), language)
        language_element.set(qn("w:eastAsia"), language)
        language_element.set(qn("w:bidi"), language)


def list_number_abstract_id(doc) -> int:
    from docx.oxml.ns import qn

    style_id = doc.styles["List Number"].style_id
    numbering = doc.part.numbering_part.element
    for abstract in numbering.findall(qn("w:abstractNum")):
        for paragraph_style in abstract.findall(f".//{qn('w:pStyle')}"):
            if paragraph_style.get(qn("w:val")) == style_id:
                return int(abstract.get(qn("w:abstractNumId")))
    raise RuntimeError("List Number abstract numbering definition not found")


def fresh_numbering_id(doc, abstract_id: int) -> int:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    numbering = doc.part.numbering_part.element
    existing = [
        int(element.get(qn("w:numId")))
        for element in numbering.findall(qn("w:num"))
    ]
    num_id = max(existing, default=0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), str(abstract_id))
    num.append(abstract)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), "1")
    override.append(start)
    num.append(override)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id: int):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    properties = paragraph._p.get_or_add_pPr()
    num_properties = properties.find(qn("w:numPr"))
    if num_properties is None:
        num_properties = OxmlElement("w:numPr")
        properties.append(num_properties)
    level = OxmlElement("w:ilvl")
    level.set(qn("w:val"), "0")
    number = OxmlElement("w:numId")
    number.set(qn("w:val"), str(num_id))
    num_properties.append(level)
    num_properties.append(number)


def build_docx(guide_md: str, output: Path):
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    doc.core_properties.language = "en-CA"
    doc.core_properties.title = "Tender Starter Guide"
    doc.core_properties.created = FIXED_INSTANT.replace(tzinfo=None)
    doc.core_properties.modified = FIXED_INSTANT.replace(tzinfo=None)
    set_style_language(doc, "en-CA")

    abstract_id = list_number_abstract_id(doc)
    lines = guide_md.split("\n")
    first_heading = True
    numbering_id = None
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("> "):
            numbering_id = None
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            run = p.add_run(stripped[2:].strip())
            run.italic = True
            run.font.size = Pt(11)
        elif stripped.startswith("## "):
            numbering_id = None
            doc.add_heading(stripped[3:].strip(), level=2)
            if first_heading:
                first_heading = False
        elif stripped.startswith("# "):
            numbering_id = None
            doc.add_heading(stripped[2:].strip(), level=1)
            if first_heading:
                first_heading = False
        elif stripped.startswith("---"):
            numbering_id = None
            continue
        elif stripped.startswith("- ") or stripped.startswith("* "):
            numbering_id = None
            doc.add_paragraph(stripped[2:].strip(), style="List Bullet")
        elif stripped[0:2].rstrip(".").isdigit() and "." in stripped[:4]:
            # Numbered list item.
            parts = stripped.split(". ", 1)
            if len(parts) == 2:
                if numbering_id is None:
                    numbering_id = fresh_numbering_id(doc, abstract_id)
                para = doc.add_paragraph(parts[1].strip(), style="List Number")
                apply_numbering(para, numbering_id)
            else:
                numbering_id = None
                doc.add_paragraph(stripped)
        else:
            numbering_id = None
            doc.add_paragraph(stripped)

    doc.save(str(output))


# --- PDF/UA export ----------------------------------------------------------

def export_pdfua(docx_path: Path, pdf_path: Path, soffice: str):
    """Export a tagged PDF/UA-compliant PDF via LibreOffice Writer."""
    filter_data = (
        '{"PDFUACompliance":{"type":"boolean","value":"true"},'
        '"UseTaggedPDF":{"type":"boolean","value":"true"},'
        '"ExportBookmarks":{"type":"boolean","value":"true"}}'
    )
    # Force a clean, temporary user profile for each invocation so that
    # accumulated LibreOffice state cannot make builds non-reproducible.
    profile_dir = pdf_path.parent / "lo-profile"
    cmd = [
        soffice,
        f"-env:UserInstallation=file://{profile_dir}",
        "--headless",
        "--norestore",
        "--nolockcheck",
        "--convert-to",
        f'pdf:writer_pdf_Export:{filter_data}',
        "--outdir",
        str(pdf_path.parent),
        str(docx_path),
    ]
    export_environment = os.environ.copy()
    export_environment["SOURCE_DATE_EPOCH"] = str(FIXED_EPOCH)
    result = subprocess.run(
        cmd,
        capture_output=True,
        text=True,
        timeout=120,
        env=export_environment,
    )
    if result.returncode != 0:
        raise RuntimeError(f"LibreOffice PDF export failed:\n{result.stderr}")

    # LibreOffice names the output after the input; rename if needed.
    produced = pdf_path.parent / (docx_path.stem + ".pdf")
    if produced != pdf_path and produced.exists():
        produced.rename(pdf_path)

    if not pdf_path.exists():
        raise RuntimeError("PDF was not produced")


# --- PDF normalization ------------------------------------------------------

def normalize_pdf(pdf_path: Path):
    from pypdf import PdfReader, PdfWriter

    reader = PdfReader(str(pdf_path))
    writer = PdfWriter(clone_from=reader)
    xmp = writer.xmp_metadata
    if xmp is None:
        raise RuntimeError("Cannot normalize PDF without XMP metadata")

    original_xmp = xmp.stream.get_data()
    normalized_xmp, replacement_count = re.subn(
        rb"\d{4}-\d{2}-\d{2}T\d{2}:\d{2}:\d{2}(?:Z|[+-]\d{2}:\d{2})",
        FIXED_XMP_DATE,
        original_xmp,
    )
    if replacement_count == 0:
        raise RuntimeError("Expected LibreOffice XMP timestamps were not found")
    writer.xmp_metadata = normalized_xmp

    writer.metadata = None
    writer.add_metadata(
        {
            "/Title": "Tender Starter Guide",
            "/Author": "KinNoKi Labs",
            "/Creator": "KinNoKi Labs deterministic tender-pack builder",
            "/Producer": "LibreOffice normalized by pypdf",
            "/CreationDate": FIXED_PDF_DATE,
            "/ModDate": FIXED_PDF_DATE,
        }
    )

    # Enforce the exact catalog /Lang — LibreOffice derives it from the UI
    # locale, which may differ from the document language setting.
    from pypdf.generic import NameObject, TextStringObject
    writer._root_object[NameObject("/Lang")] = TextStringObject("en-CA")

    # The cloned LibreOffice file carries its old XMP object and trailer ID.
    # Remove unreachable objects, then set a deterministic ID derived from the
    # normalized content (generate_file_identifiers uses randomness, which
    # breaks reproducibility).
    writer.compress_identical_objects(
        remove_duplicates=True,
        remove_unreferenced=True,
    )
    writer._ID = None  # pypdf 6.10 has no public API to clear a cloned trailer ID.

    # Serialize once without an /ID to get a stable content digest.
    temporary_for_hash = pdf_path.with_suffix(".for-hash.pdf")
    writer.write(temporary_for_hash)
    digest = hashlib.md5(temporary_for_hash.read_bytes()).digest()

    # Set a deterministic /ID and write the final normalized PDF.
    from pypdf.generic import ArrayObject, ByteStringObject
    writer._ID = ArrayObject(
        [ByteStringObject(digest), ByteStringObject(digest)]
    )

    temporary = pdf_path.with_suffix(".normalized.pdf")
    writer.write(temporary)
    temporary_for_hash.unlink(missing_ok=True)
    temporary.replace(pdf_path)


# --- PDF verification -------------------------------------------------------

def verify_pdf(pdf_path: Path):
    from pypdf import PdfReader

    reader = PdfReader(str(pdf_path))
    raw = pdf_path.read_bytes()
    head = raw[:4096].decode("latin1")
    full = raw.decode("latin1")

    errors = []
    if not head.startswith("%PDF"):
        errors.append("missing %PDF header")
    if "/StructTreeRoot" not in full:
        errors.append("missing /StructTreeRoot (tagged PDF)")
    if "/MarkInfo" not in full:
        errors.append("missing /MarkInfo")

    # Check title.
    meta = reader.metadata
    if not meta or not (meta.title or "").strip():
        errors.append("missing document title")

    # Exact catalog /Lang must equal en-CA.
    catalog_language = str(reader.trailer["/Root"].get("/Lang", ""))
    if catalog_language != "en-CA":
        errors.append(f"catalog /Lang must be en-CA, found {catalog_language or 'missing'}")

    # XMP packet must carry the PDF/UA declaration and normalized timestamps.
    xmp_bytes = reader.xmp_metadata.stream.get_data() if reader.xmp_metadata else b""
    if b"pdfuaid:part" not in xmp_bytes:
        errors.append("missing pdfuaid:part (PDF/UA marker)")
    if FIXED_XMP_DATE not in xmp_bytes:
        errors.append("XMP timestamps were not normalized")

    # Verify independent numbered lists both begin at 1.
    extracted = "\n".join(page.extract_text() or "" for page in reader.pages)
    if "1. Open the official notice" not in extracted:
        errors.append("review-method list must start at 1")
    if "1. Recheck the official notice" not in extracted:
        errors.append("next-steps list must restart at 1")
    if "6. Recheck the official notice" in extracted:
        errors.append("next-steps list incorrectly continues at 6")

    if errors:
        raise RuntimeError(f"PDF structure verification failed: {'; '.join(errors)}")

    print(f"PDF verified: {len(reader.pages)} page(s), title='{meta.title}', /Lang={catalog_language}")


# --- XLSX -------------------------------------------------------------------

def build_xlsx(data: dict, output: Path):
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Font, PatternFill
    from openpyxl.utils import get_column_letter

    wb = Workbook()
    # Remove default sheet.
    wb.remove(wb.active)

    header_font = Font(name="Arial", bold=True, size=11, color="FFFFFF")
    header_fill = PatternFill(start_color="2C2C2E", end_color="2C2C2E", fill_type="solid")
    header_align = Alignment(horizontal="left", vertical="center", wrap_text=True)
    cell_align = Alignment(horizontal="left", vertical="top", wrap_text=True)
    body_font = Font(name="Arial", size=11)

    for sheet_def in data["sheets"]:
        ws = wb.create_sheet(title=sheet_def["name"])
        headers = sheet_def["headers"]
        rows = sheet_def.get("rows", [])

        # Write headers.
        for col_idx, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col_idx, value=header)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = header_align

        # Write data rows.
        for row_idx, row_data in enumerate(rows, 2):
            for col_idx, value in enumerate(row_data, 1):
                cell = ws.cell(row=row_idx, column=col_idx, value=value)
                cell.font = body_font
                cell.alignment = cell_align

        # Freeze header row.
        ws.freeze_panes = "A2"

        # Enable autofilter on header row.
        if rows:
            last_col = get_column_letter(len(headers))
            ws.auto_filter.ref = f"A1:{last_col}{len(rows) + 1}"

        # Set practical column widths.
        for col_idx, header in enumerate(headers, 1):
            max_len = len(header)
            for row_data in rows:
                if col_idx <= len(row_data) and row_data[col_idx - 1]:
                    max_len = max(max_len, min(len(str(row_data[col_idx - 1])), 50))
            ws.column_dimensions[get_column_letter(col_idx)].width = min(max(max_len + 4, 12), 55)

    wb.properties.created = FIXED_INSTANT.replace(tzinfo=None)
    wb.properties.modified = FIXED_INSTANT.replace(tzinfo=None)
    wb.save(str(output))


# --- Deterministic ZIP ------------------------------------------------------

def create_deterministic_zip(output: Path, entries: list[tuple[str, bytes]]):
    """Create a ZIP with fixed entry timestamps and no compression-level variance."""
    output.parent.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(str(output), "w", compression=zipfile.ZIP_DEFLATED, compresslevel=9) as zf:
        for name, data in entries:
            # Reset all entry metadata to fixed values for reproducibility.
            info = zipfile.ZipInfo(filename=name, date_time=FIXED_TIMESTAMP)
            info.compress_type = zipfile.ZIP_DEFLATED
            info.external_attr = 0o644 << 16  # regular file permissions
            info.create_system = 3  # Unix
            zf.writestr(info, data)


# --- Office-container normalization -----------------------------------------

def normalize_zip_container(path: Path):
    original = path.read_bytes()
    with tempfile.TemporaryDirectory(prefix="normalize-office-") as directory:
        source = Path(directory) / "source.zip"
        normalized = Path(directory) / "normalized.zip"
        source.write_bytes(original)
        with zipfile.ZipFile(source, "r") as reader:
            entries = [(info.filename, reader.read(info.filename)) for info in reader.infolist()]

        # Normalize docProps/core.xml timestamps (openpyxl overwrites modified at save).
        normalized_entries = []
        for name, data in entries:
            if name == "docProps/core.xml":
                data = re.sub(
                    rb"<dcterms:modified[^>]*>[^<]*</dcterms:modified>",
                    b'<dcterms:modified xmlns:dcterms="http://purl.org/dc/terms/" '
                    b'xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance" '
                    b'xsi:type="dcterms:W3CDTF">'
                    + FIXED_XMP_DATE +
                    b"</dcterms:modified>",
                    data,
                )
            normalized_entries.append((name, data))

        create_deterministic_zip(normalized, sorted(normalized_entries))
        path.write_bytes(normalized.read_bytes())


if __name__ == "__main__":
    main()
